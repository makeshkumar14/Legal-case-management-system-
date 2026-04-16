from __future__ import annotations

import argparse
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
DEFAULT_TEMPLATE = Path(r"c:\Users\MAKESH\OneDrive\Desktop\MAX\DBMS REPORT .docx")
DEFAULT_OUTPUT = ROOT / "Legal_Case_Management_DBMS_Report.docx"
ASSET_DIR = ROOT / "report_assets"


ABBREVIATIONS = [
    "API : Application Programming Interface",
    "DBMS : Database Management System",
    "ER : Entity Relationship",
    "JWT : JSON Web Token",
    "LCMS : Legal Case Management System",
    "ORM : Object Relational Mapper",
    "OTP : One Time Password",
    "RDBMS : Relational Database Management System",
    "SQL : Structured Query Language",
    "UI : User Interface",
]


FIGURES = [
    ("Fig. 3.1", "ER Diagram of LCMS"),
    ("Fig. 3.2", "Use Case Diagram of LCMS"),
    ("Fig. 3.3", "Architecture Diagram of LCMS"),
    ("Fig. 5.1", "Role-Based Dashboard Overview"),
    ("Fig. 5.2", "Case Processing Workflow"),
]


REFERENCES = [
    "Project source files: backend/app.py, backend/schema.sql, backend/routes/*.py, frontend/src/App.jsx, frontend/src/services/api.js.",
    "Flask Documentation. Pallets Projects. https://flask.palletsprojects.com/",
    "SQLAlchemy Documentation. https://docs.sqlalchemy.org/",
    "React Documentation. https://react.dev/",
    "Vite Guide. https://vite.dev/guide/",
    "MySQL Reference Manual. https://dev.mysql.com/doc/",
    "Flask-JWT-Extended Documentation. https://flask-jwt-extended.readthedocs.io/",
]


TABLE_DATA = [
    ("users", "Stores citizens, advocates, and court admins with role-specific identifiers."),
    ("otp_codes", "Keeps time-bound OTPs for citizen authentication."),
    ("courtrooms", "Tracks room status, judge, and active courtroom details."),
    ("cases", "Maintains case metadata, parties, advocate assignment, hearing date, and status."),
    ("hearings", "Stores hearing schedule, type, timing, notes, and location."),
    ("case_timeline", "Captures major events in each case lifecycle."),
    ("documents", "Stores evidence and case file metadata, uploader, size, and verification status."),
    ("tasks", "Tracks advocate task items linked to a case."),
    ("case_notes", "Stores internal notes made by advocates or court staff."),
    ("notifications", "Keeps alerts for hearings, updates, and reminders."),
    ("messages", "Stores direct communication between users."),
]


OBJECTIVES = [
    "To design a centralized Legal Case Management System that replaces fragmented and manual case handling.",
    "To provide secure role-based access for citizens, advocates, and court administrators.",
    "To maintain accurate case, hearing, document, note, and task records in a relational database.",
    "To improve transparency by allowing citizens to track their cases and upcoming hearings digitally.",
    "To support advocates with evidence management, case notes, pending tasks, and performance analytics.",
    "To help court administrators monitor pendency, hearing schedules, courtroom activity, and advocate workloads.",
    "To ensure data consistency and controlled access through MySQL, SQLAlchemy relationships, JWT authentication, and validation logic.",
]


EXISTING_LIMITATIONS = [
    "Case records are often distributed across paper files, spreadsheets, and disconnected software, making search and retrieval slow.",
    "Citizens have limited visibility into the latest status of their cases and must rely on manual follow-up.",
    "Hearing schedules, case notes, and evidence tracking are difficult to coordinate when different stakeholders maintain separate records.",
    "Lack of centralized notifications increases the chance of missed hearings, delayed submissions, and inconsistent communication.",
    "Manual workflows introduce duplication, update delays, and errors in status tracking, advocate assignment, and courtroom planning.",
]


MODULES = [
    ("Authentication and Identity", "Supports citizen OTP login, advocate credential login, court admin login, and shared profile management."),
    ("Case Management", "Creates, updates, filters, and searches cases with role-based visibility and QR-based lookup support."),
    ("Hearing Scheduler", "Schedules hearings, updates next-hearing details, and exposes calendar-friendly event data."),
    ("Document Management", "Uploads evidence and legal files, stores metadata, and supports verification workflows."),
    ("Task and Notes Module", "Lets advocates and staff maintain case-specific tasks and structured notes."),
    ("Messaging and Notifications", "Delivers in-app notifications, read/unread states, direct user messaging, and optional email alerts."),
    ("Courtroom and Analytics", "Shows courtroom status, dashboard summaries, pendency, case mix, hearing trends, and advocate performance."),
]


FUNCTIONAL_RESULTS = [
    "Citizens can register with Aadhaar details, receive a demo OTP, verify the OTP, and view only their own cases.",
    "Advocates can log in with a Bar Council ID, view assigned cases, manage documents, tasks, notes, and messaging data.",
    "Court admins can view all cases, create and update case records, schedule hearings, manage courtroom data, and monitor system analytics.",
    "Updating a case can automatically create hearing entries, timeline events, and notifications for relevant users.",
    "The application exposes filtered dashboard metrics such as total cases, pending matters, hearing counts, evidence totals, and advocate performance.",
]


FUTURE_ENHANCEMENTS = [
    "Replace demo OTP responses with real SMS gateway integration.",
    "Add production-grade audit logs for every critical database change.",
    "Introduce document versioning and cloud object storage for large evidence files.",
    "Add judge-wise scheduling conflict detection and courtroom slot optimization.",
    "Support e-filing workflows, digital signatures, and PDF report export.",
    "Strengthen production security with refresh tokens, stronger secret management, and role-based permission matrices.",
]


def get_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    candidates = []
    if bold:
        candidates.extend(["arialbd.ttf", "calibrib.ttf"])
    else:
        candidates.extend(["arial.ttf", "calibri.ttf"])

    for candidate in candidates:
        try:
            return ImageFont.truetype(candidate, size)
        except OSError:
            continue
    return ImageFont.load_default()


def clear_document(document: Document) -> None:
    body = document._body._element
    for child in list(body):
        if child.tag.endswith("sectPr"):
            continue
        body.remove(child)


def resolve_style(document: Document, style: str | None) -> str | None:
    if not style:
        return None
    try:
        document.styles[style]
        return style
    except KeyError:
        for candidate in document.styles:
            if getattr(candidate, "name", None) == style:
                return candidate
    return style


def add_paragraph(document: Document, text: str = "", style: str = "Body Text", align=None):
    paragraph = document.add_paragraph(style=resolve_style(document, style))
    if align is not None:
        paragraph.alignment = align
    if text:
        paragraph.add_run(text)
    return paragraph


def add_body(document: Document, text: str) -> None:
    for chunk in [part.strip() for part in text.split("\n\n") if part.strip()]:
        add_paragraph(document, chunk, style="Body Text")


def add_list(document: Document, items: list[str]) -> None:
    for item in items:
        add_paragraph(document, item, style="List Paragraph")


def add_title(document: Document, text: str, size: int = 18) -> None:
    p = add_paragraph(document, style="Normal", align=WD_ALIGN_PARAGRAPH.CENTER)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)


def add_cover_page(document: Document) -> None:
    add_title(document, "LEGAL CASE MANAGEMENT SYSTEM", 20)
    add_paragraph(document, "", style="Body Text")
    add_paragraph(document, "Submitted by", style="Body Text", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(
        document,
        "[Student Name 1]  [Register No.]\n[Student Name 2]  [Register No.]\n[Student Name 3]  [Register No.]",
        style="Normal",
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    add_paragraph(document, "", style="Body Text")
    add_paragraph(document, "Under the guidance of [Guide Name]", style="Normal", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(
        document,
        "(Guide designation / Department)",
        style="Normal",
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    add_paragraph(document, "", style="Body Text")
    add_paragraph(
        document,
        "DATABASE MANAGEMENT SYSTEM PROJECT REPORT",
        style="Normal",
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    add_paragraph(document, "COMPUTER SCIENCE AND ENGINEERING", style="Normal", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(document, "FACULTY OF ENGINEERING AND TECHNOLOGY", style="Body Text", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(document, "", style="Body Text")
    add_paragraph(
        document,
        "SRM INSTITUTE OF SCIENCE AND TECHNOLOGY\nRAMAPURAM, CHENNAI",
        style="Heading 1",
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    add_paragraph(document, "APRIL 2026", style="Body Text", align=WD_ALIGN_PARAGRAPH.CENTER)
    document.add_page_break()


def add_certificate(document: Document) -> None:
    add_paragraph(document, "SRM INSTITUTE OF SCIENCE AND TECHNOLOGY", style="Normal", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(document, "(Deemed to be University U/S 3 of UGC Act, 1956)", style="Normal", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(document, "", style="Body Text")
    add_title(document, "BONAFIDE CERTIFICATE", 16)
    add_body(
        document,
        "Certified that this project report titled \"LEGAL CASE MANAGEMENT SYSTEM\" is the bonafide work of "
        "[Student Name 1], [Student Name 2], and [Student Name 3] who carried out the project work under my "
        "supervision.\n\n"
        "Certified further that, to the best of my knowledge, the work reported herein does not form part of any "
        "other project report or dissertation on the basis of which a degree or award was conferred on any other "
        "candidate."
    )
    add_paragraph(document, "", style="Body Text")
    add_paragraph(document, "SIGNATURE", style="Normal")
    add_paragraph(document, "[Guide Name]\nGuide / Faculty", style="Normal")
    add_paragraph(document, "", style="Body Text")
    add_paragraph(document, "SIGNATURE", style="Normal")
    add_paragraph(document, "[Head of Department]\nProfessor and Head", style="Normal")
    add_paragraph(document, "", style="Body Text")
    add_paragraph(
        document,
        "Submitted for the project viva-voce held on __________ at SRM Institute of Science and Technology, Ramapuram, Chennai.",
        style="Normal",
    )
    document.add_page_break()


def add_declaration(document: Document) -> None:
    add_paragraph(
        document,
        "SRM INSTITUTE OF SCIENCE AND TECHNOLOGY\nRAMAPURAM, CHENNAI",
        style="Normal",
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    add_title(document, "DECLARATION", 16)
    add_body(
        document,
        "We hereby declare that the entire work contained in this project report titled \"LEGAL CASE MANAGEMENT "
        "SYSTEM\" has been carried out by [Student Name 1], [Student Name 2], and [Student Name 3] under the "
        "guidance of [Guide Name], Department of Computer Science and Engineering.\n\n"
        "The content presented in this report is original to this project except where due acknowledgment has been "
        "made. The report has been prepared for academic submission in the Database Management System course."
    )
    add_paragraph(document, "Place : Chennai", style="Normal")
    add_paragraph(document, "Date  : __________", style="Normal")
    add_paragraph(document, "", style="Body Text")
    add_paragraph(document, "[Student Name 1]", style="Normal")
    add_paragraph(document, "[Student Name 2]", style="Normal")
    add_paragraph(document, "[Student Name 3]", style="Normal")
    document.add_page_break()


def add_abstract(document: Document) -> None:
    add_title(document, "ABSTRACT", 16)
    add_body(
        document,
        "The Legal Case Management System (LCMS) is a full-stack database-driven application developed to "
        "digitize and streamline case handling for citizens, advocates, and court administrators. The system "
        "addresses common issues in legal workflows such as fragmented records, delayed status updates, poor "
        "communication, and limited public visibility into case progress.\n\n"
        "The proposed system uses a React and Vite frontend, a Flask backend, SQLAlchemy for object-relational "
        "mapping, and MySQL as the central relational database. Core modules include citizen registration and OTP "
        "authentication, advocate and court login, case creation and tracking, hearing scheduling, document "
        "management, case notes, tasks, notifications, messaging, courtroom monitoring, and analytics dashboards.\n\n"
        "A normalized relational schema supports entities such as users, cases, hearings, timelines, documents, "
        "tasks, notes, courtrooms, notifications, and messages. Role-based access ensures that citizens see only "
        "their own matters, advocates see assigned cases, and court administrators retain broad operational "
        "control. The system improves transparency, supports timely hearing coordination, and provides a stronger "
        "foundation for future court digitization."
    )
    document.add_page_break()


def add_front_matter(document: Document) -> None:
    add_title(document, "TABLE OF CONTENTS", 16)
    toc_entries = [
        "Abstract",
        "List of Figures",
        "List of Abbreviations",
        "Chapter 1  Introduction",
        "Chapter 2  Existing System",
        "Chapter 3  Design",
        "Chapter 4  Proposed Methodology",
        "Chapter 5  Implementation",
        "Chapter 6  Result and Discussion",
        "Chapter 7  Conclusion",
        "Chapter 8  References",
    ]
    add_list(document, toc_entries)

    add_paragraph(document, "", style="Body Text")
    add_title(document, "LIST OF FIGURES", 16)
    add_list(document, [f"{prefix}  {name}" for prefix, name in FIGURES])

    add_paragraph(document, "", style="Body Text")
    add_title(document, "LIST OF ABBREVIATIONS", 16)
    add_list(document, ABBREVIATIONS)
    document.add_page_break()


def add_chapter_heading(document: Document, number: int, title: str) -> None:
    add_paragraph(document, f"CHAPTER {number}", style="Normal", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(document, title.upper(), style="Normal", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(document, "", style="Body Text")


def add_section_heading(document: Document, text: str) -> None:
    add_paragraph(document, text, style="Heading 1")


def add_table(document: Document, rows: list[tuple[str, str]], headings: tuple[str, str]) -> None:
    table = document.add_table(rows=1, cols=2)
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    header = table.rows[0].cells
    header[0].text = headings[0]
    header[1].text = headings[1]

    for left, right in rows:
        cells = table.add_row().cells
        cells[0].text = left
        cells[1].text = right


def draw_box(draw: ImageDraw.ImageDraw, xy: tuple[int, int, int, int], title: str, lines: list[str], fill: str) -> None:
    draw.rounded_rectangle(xy, radius=18, fill=fill, outline="#1f2937", width=3)
    title_font = get_font(30, bold=True)
    body_font = get_font(22)
    x1, y1, _, _ = xy
    draw.text((x1 + 20, y1 + 18), title, font=title_font, fill="#0f172a")
    text_y = y1 + 70
    for line in lines:
        draw.text((x1 + 20, text_y), line, font=body_font, fill="#111827")
        text_y += 32


def draw_arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], label: str | None = None) -> None:
    draw.line([start, end], fill="#0f172a", width=5)
    ex, ey = end
    sx, sy = start
    if abs(ex - sx) >= abs(ey - sy):
        if ex >= sx:
            points = [(ex, ey), (ex - 16, ey - 8), (ex - 16, ey + 8)]
        else:
            points = [(ex, ey), (ex + 16, ey - 8), (ex + 16, ey + 8)]
    else:
        if ey >= sy:
            points = [(ex, ey), (ex - 8, ey - 16), (ex + 8, ey - 16)]
        else:
            points = [(ex, ey), (ex - 8, ey + 16), (ex + 8, ey + 16)]
    draw.polygon(points, fill="#0f172a")
    if label:
        font = get_font(20, bold=True)
        mid_x = (sx + ex) // 2
        mid_y = (sy + ey) // 2
        draw.text((mid_x + 6, mid_y - 20), label, font=font, fill="#0f172a")


def create_er_diagram(path: Path) -> None:
    image = Image.new("RGB", (1800, 1200), "#f8fafc")
    draw = ImageDraw.Draw(image)
    title_font = get_font(40, bold=True)
    draw.text((600, 30), "LCMS Entity Relationship Diagram", font=title_font, fill="#0f172a")

    draw_box(draw, (60, 130, 470, 360), "Users", ["id (PK)", "name, role, phone", "aadhaar/bar/admin ids"], "#dbeafe")
    draw_box(draw, (630, 110, 1110, 380), "Cases", ["id (PK), case_number", "petitioner, respondent", "status, priority, advocate_id", "courtroom_id, next_hearing"], "#dcfce7")
    draw_box(draw, (1260, 130, 1710, 340), "Courtrooms", ["id (PK)", "name, judge, status", "current_case"], "#fde68a")

    draw_box(draw, (110, 500, 480, 690), "OTP Codes", ["id (PK)", "aadhaar, otp_code", "expires_at, used"], "#fae8ff")
    draw_box(draw, (560, 500, 900, 730), "Hearings", ["id (PK), case_id (FK)", "date, type, location", "start_time, end_time"], "#ffe4e6")
    draw_box(draw, (960, 500, 1320, 730), "Documents", ["id (PK), case_id (FK)", "uploaded_by (FK)", "title, file_type, verified"], "#fef3c7")
    draw_box(draw, (1380, 500, 1710, 730), "Tasks", ["id (PK), case_id (FK)", "user_id (FK)", "title, due_date"], "#e0f2fe")

    draw_box(draw, (260, 840, 630, 1060), "Case Timeline", ["id (PK), case_id (FK)", "date, event, description"], "#ede9fe")
    draw_box(draw, (760, 840, 1120, 1060), "Case Notes", ["id (PK), case_id (FK)", "user_id (FK), content"], "#fce7f3")
    draw_box(draw, (1220, 840, 1560, 1060), "Notifications", ["id (PK), user_id (FK)", "type, title, is_read"], "#d9f99d")
    draw_box(draw, (1460, 830, 1760, 1060), "Messages", ["id (PK)", "sender_id, receiver_id", "content, sent_at"], "#fee2e2")

    draw_arrow(draw, (470, 240), (630, 240), "1:M")
    draw_arrow(draw, (1110, 240), (1260, 240), "M:1")
    draw_arrow(draw, (260, 360), (260, 500), "1:M")
    draw_arrow(draw, (760, 380), (720, 500), "1:M")
    draw_arrow(draw, (950, 380), (1080, 500), "1:M")
    draw_arrow(draw, (1080, 380), (1540, 500), "1:M")
    draw_arrow(draw, (760, 380), (450, 840), "1:M")
    draw_arrow(draw, (840, 380), (940, 840), "1:M")
    draw_arrow(draw, (320, 360), (1350, 840), "1:M")
    draw_arrow(draw, (360, 360), (1600, 840), "1:M")

    image.save(path)


def create_use_case_diagram(path: Path) -> None:
    image = Image.new("RGB", (1800, 1100), "#fffdf8")
    draw = ImageDraw.Draw(image)
    title_font = get_font(40, bold=True)
    draw.text((620, 24), "LCMS Use Case Diagram", font=title_font, fill="#0f172a")

    actor_font = get_font(28, bold=True)
    use_case_font = get_font(24)
    ellipse_fill = "#eef2ff"

    def actor(x: int, y: int, name: str) -> None:
        draw.ellipse((x - 35, y, x + 35, y + 70), outline="#111827", width=4)
        draw.line((x, y + 70, x, y + 180), fill="#111827", width=4)
        draw.line((x - 60, y + 105, x + 60, y + 105), fill="#111827", width=4)
        draw.line((x, y + 180, x - 50, y + 260), fill="#111827", width=4)
        draw.line((x, y + 180, x + 50, y + 260), fill="#111827", width=4)
        draw.text((x - 70, y + 280), name, font=actor_font, fill="#111827")

    def use_case(x1: int, y1: int, x2: int, y2: int, text: str) -> None:
        draw.ellipse((x1, y1, x2, y2), fill=ellipse_fill, outline="#1f2937", width=3)
        draw.multiline_text((x1 + 28, y1 + 28), text, font=use_case_font, fill="#111827", spacing=4)

    actor(170, 180, "Citizen")
    actor(170, 520, "Advocate")
    actor(170, 840, "Court Admin")

    use_case(500, 120, 930, 220, "Register / Login")
    use_case(500, 260, 930, 360, "View Case Status\nand Hearing Details")
    use_case(500, 400, 930, 500, "Upload / Review\nDocuments")
    use_case(500, 540, 930, 640, "Manage Case Notes\nand Tasks")
    use_case(500, 680, 930, 780, "Send Messages and\nReceive Notifications")
    use_case(500, 820, 930, 920, "Schedule Hearings /\nUpdate Case Status")
    use_case(500, 960, 930, 1060, "View Analytics /\nCourtroom Dashboard")

    for y in [170, 310, 450, 590, 730, 870, 1010]:
        draw_arrow(draw, (250, 250), (500, y))
    for y in [450, 590, 730]:
        draw_arrow(draw, (250, 590), (500, y))
    for y in [730, 870, 1010]:
        draw_arrow(draw, (250, 910), (500, y))

    image.save(path)


def create_architecture_diagram(path: Path) -> None:
    image = Image.new("RGB", (1800, 1000), "#f8fafc")
    draw = ImageDraw.Draw(image)
    title_font = get_font(40, bold=True)
    draw.text((550, 30), "LCMS System Architecture", font=title_font, fill="#0f172a")

    draw_box(draw, (70, 180, 460, 480), "Users", ["Citizens", "Advocates", "Court Admins"], "#e0f2fe")
    draw_box(draw, (600, 150, 1120, 520), "React Frontend", ["Role-based routing", "Dashboards and pages", "Context + API service layer"], "#dcfce7")
    draw_box(draw, (1280, 150, 1720, 520), "Flask API", ["Blueprint routes", "JWT auth", "Validation and business logic"], "#fee2e2")
    draw_box(draw, (620, 640, 1080, 910), "SQLAlchemy ORM", ["Models and relationships", "CRUD operations", "Transaction handling"], "#ede9fe")
    draw_box(draw, (1270, 640, 1710, 910), "MySQL Database", ["users, cases, hearings", "documents, tasks, messages", "notifications, notes, timelines"], "#fde68a")
    draw_box(draw, (70, 640, 470, 910), "Supporting Services", ["Uploads folder", "Flask-Mail", "OTP generation", "Analytics endpoints"], "#fef3c7")

    draw_arrow(draw, (460, 330), (600, 330), "HTTPS")
    draw_arrow(draw, (1120, 330), (1280, 330), "REST API")
    draw_arrow(draw, (1490, 520), (1490, 640), "ORM")
    draw_arrow(draw, (1080, 780), (1270, 780), "SQL")
    draw_arrow(draw, (470, 770), (620, 770), "Files / Mail / OTP")

    image.save(path)


def create_dashboard_figure(path: Path) -> None:
    image = Image.new("RGB", (1800, 1000), "#fffaf5")
    draw = ImageDraw.Draw(image)
    title_font = get_font(40, bold=True)
    draw.text((500, 26), "Role-Based Dashboard Overview", font=title_font, fill="#0f172a")

    draw_box(draw, (80, 140, 560, 860), "Citizen Dashboard", ["My cases", "Next hearing", "Case documents", "OTP-based login"], "#dbeafe")
    draw_box(draw, (660, 140, 1140, 860), "Advocate Dashboard", ["Assigned cases", "Pending tasks", "Case notes", "Performance summary"], "#dcfce7")
    draw_box(draw, (1240, 140, 1720, 860), "Court Dashboard", ["Total / pending cases", "Today's hearings", "Courtroom board", "Analytics and advocate stats"], "#fee2e2")

    stat_font = get_font(22, bold=True)
    draw.rounded_rectangle((150, 520, 490, 610), radius=16, fill="#93c5fd", outline="#1e3a8a", width=2)
    draw.text((185, 550), "Transparent citizen case tracking", font=stat_font, fill="#0f172a")
    draw.rounded_rectangle((730, 520, 1070, 610), radius=16, fill="#86efac", outline="#166534", width=2)
    draw.text((792, 550), "Advocate workflow support", font=stat_font, fill="#0f172a")
    draw.rounded_rectangle((1310, 520, 1650, 610), radius=16, fill="#fca5a5", outline="#991b1b", width=2)
    draw.text((1382, 550), "Court-level monitoring", font=stat_font, fill="#0f172a")

    image.save(path)


def create_workflow_figure(path: Path) -> None:
    image = Image.new("RGB", (1800, 900), "#f8fafc")
    draw = ImageDraw.Draw(image)
    title_font = get_font(40, bold=True)
    draw.text((560, 28), "LCMS Case Processing Workflow", font=title_font, fill="#0f172a")

    steps = [
        "Citizen / Advocate\nLogin",
        "Case Creation or\nRegistration",
        "Advocate Assignment\nand Review",
        "Hearing Scheduling\nand Calendar Update",
        "Documents, Notes,\nTasks, Messaging",
        "Status Update,\nNotifications, Analytics",
        "Closure /\nArchival",
    ]
    fills = ["#dbeafe", "#c7d2fe", "#dcfce7", "#fde68a", "#fbcfe8", "#fecaca", "#e5e7eb"]
    x = 60
    for index, step in enumerate(steps):
        draw.rounded_rectangle((x, 320, x + 220, 520), radius=22, fill=fills[index], outline="#1f2937", width=3)
        draw.multiline_text((x + 24, 374), step, font=get_font(24, bold=True), fill="#111827", spacing=6)
        if index < len(steps) - 1:
            draw_arrow(draw, (x + 220, 420), (x + 280, 420))
        x += 250

    image.save(path)


def build_assets() -> dict[str, Path]:
    ASSET_DIR.mkdir(exist_ok=True)
    assets = {
        "er": ASSET_DIR / "fig_3_1_er_diagram.png",
        "use_case": ASSET_DIR / "fig_3_2_use_case.png",
        "architecture": ASSET_DIR / "fig_3_3_architecture.png",
        "dashboards": ASSET_DIR / "fig_5_1_dashboards.png",
        "workflow": ASSET_DIR / "fig_5_2_workflow.png",
    }

    create_er_diagram(assets["er"])
    create_use_case_diagram(assets["use_case"])
    create_architecture_diagram(assets["architecture"])
    create_dashboard_figure(assets["dashboards"])
    create_workflow_figure(assets["workflow"])
    return assets


def add_figure(document: Document, image_path: Path, caption: str, width: float = 6.5) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(width))

    cap = add_paragraph(document, caption, style="Normal", align=WD_ALIGN_PARAGRAPH.CENTER)
    if cap.runs:
        cap.runs[0].italic = True


def chapter_one(document: Document) -> None:
    add_chapter_heading(document, 1, "Introduction")

    add_section_heading(document, "Introduction:")
    add_body(
        document,
        "Legal information systems generate a large amount of structured and semi-structured data, including user "
        "profiles, case files, hearing schedules, evidence documents, communication records, and performance "
        "statistics. A DBMS-based solution is well suited to this environment because the data has clear entities, "
        "relationships, integrity constraints, and role-sensitive access requirements.\n\n"
        "The Legal Case Management System is designed as a multi-user web platform that centralizes these records "
        "and exposes them through tailored dashboards. Citizens can track their own matters, advocates can manage "
        "assigned cases, and court administrators can control case flow, hearings, and analytics from a single system."
    )

    add_section_heading(document, "Problem Statement:")
    add_body(
        document,
        "Traditional legal record management is often slow, opaque, and fragmented. Case progress may be tracked in "
        "separate registers, spreadsheets, or disconnected software tools, making it difficult to provide accurate "
        "real-time information to all stakeholders.\n\n"
        "These gaps lead to delays in hearing communication, poor coordination among parties, inconsistent updates, "
        "and reduced transparency for citizens. The project therefore focuses on building a secure relational system "
        "that manages legal case data end-to-end and exposes it through controlled, role-based workflows."
    )

    add_section_heading(document, "Objective:")
    add_list(document, OBJECTIVES)

    add_section_heading(document, "Scope and Motivation:")
    add_body(
        document,
        "The scope of LCMS includes authentication, case registration and tracking, hearing management, document "
        "handling, case notes, tasks, messaging, notifications, courtroom monitoring, and analytics. The system is "
        "implemented as a full-stack academic prototype and is intended to demonstrate how structured database design "
        "can support a realistic legal workflow.\n\n"
        "The motivation behind the project is to improve efficiency, visibility, and coordination in court-related "
        "processes while maintaining data integrity. The system also highlights how a centralized DBMS can reduce "
        "redundancy, make updates traceable, and improve service quality for end users."
    )
    document.add_page_break()


def chapter_two(document: Document) -> None:
    add_chapter_heading(document, 2, "Existing System")
    add_body(
        document,
        "In many legal environments, a fully integrated digital workflow is absent. Records may be handled through "
        "manual registers, ad hoc spreadsheets, isolated communication channels, and separate document stores. Even "
        "when partial software support exists, it frequently fails to provide unified access to citizens, advocates, "
        "and court staff."
    )
    add_list(document, EXISTING_LIMITATIONS)
    add_body(
        document,
        "These drawbacks justify the development of a centralized database-driven application. A modern legal case "
        "management system should support both operational control and public-facing transparency while preserving "
        "role separation and relational consistency."
    )
    document.add_page_break()


def chapter_three(document: Document, assets: dict[str, Path]) -> None:
    add_chapter_heading(document, 3, "Design")

    add_section_heading(document, "Entity-Relationship (ER) Design:")
    add_body(
        document,
        "The database schema is organized around users, cases, hearings, evidence records, notes, tasks, messages, "
        "notifications, and courtroom metadata. The design uses relational keys to connect legal workflows while "
        "preserving referential integrity across tables."
    )
    add_table(document, TABLE_DATA, ("Table / Entity", "Purpose"))
    add_paragraph(document, "", style="Body Text")
    add_body(
        document,
        "Important relationships include one-to-many links between advocates and cases, cases and hearings, cases and "
        "documents, cases and tasks, cases and case notes, and users and notifications. Messages create a user-to-user "
        "communication channel, while OTP codes support citizen login through Aadhaar-linked verification."
    )
    add_figure(document, assets["er"], "Fig. 3.1 ER Diagram of LCMS")

    add_section_heading(document, "Use Case Design:")
    add_body(
        document,
        "The core actors are Citizen, Advocate, and Court Admin. Citizens primarily authenticate, track their cases, "
        "view hearings, and access documents. Advocates manage assigned matters, collaborate through notes and "
        "messaging, and monitor pending work. Court admins maintain overall operational control by registering cases, "
        "scheduling hearings, updating statuses, reviewing courtrooms, and checking analytics."
    )
    add_figure(document, assets["use_case"], "Fig. 3.2 Use Case Diagram of LCMS")

    add_section_heading(document, "System Architecture:")
    add_body(
        document,
        "The application follows a layered web architecture. The React frontend handles presentation, navigation, "
        "and API interaction. The Flask backend exposes REST endpoints and contains the business rules. SQLAlchemy "
        "maps database entities into Python models, and MySQL stores the persistent relational data. Supporting "
        "services include local file uploads, OTP generation, and Flask-Mail-based notification delivery."
    )
    add_figure(document, assets["architecture"], "Fig. 3.3 Architecture Diagram of LCMS")

    add_section_heading(document, "Front-End Design:")
    add_body(
        document,
        "The frontend uses React Router for protected role-based navigation and a shared dashboard layout for all "
        "authenticated roles. Reusable UI components such as status badges, modals, timeline views, QR code widgets, "
        "toasts, and loading skeletons create a consistent interface across the portal.\n\n"
        "The design emphasizes quick status visibility, filtering, and multi-role usability. Citizens see their own "
        "cases and updates, advocates see case workload and evidence-related tasks, and court admins receive broader "
        "operational analytics and scheduling controls."
    )
    document.add_page_break()


def chapter_four(document: Document) -> None:
    add_chapter_heading(document, 4, "Proposed Methodology")

    add_section_heading(document, "Modules Description:")
    module_rows = [(name, desc) for name, desc in MODULES]
    add_table(document, module_rows, ("Module", "Description"))

    add_paragraph(document, "", style="Body Text")
    add_section_heading(document, "Database Connectivity:")
    add_body(
        document,
        "The backend uses Flask SQLAlchemy to connect the application layer to the MySQL database defined in the "
        "configuration module. Application startup initializes the database extension and creates required tables. "
        "All core data access operations are executed through ORM models and route handlers rather than hard-coded "
        "SQL in the UI layer.\n\n"
        "JWT tokens are attached to protected requests, allowing the backend to identify the active user and apply "
        "role-based filtering. For example, citizen case listings are filtered by petitioner identity, advocate case "
        "listings are filtered by advocate assignment, and court users can access broader operational data.\n\n"
        "The methodology also includes relational event creation during business operations. When a hearing is "
        "scheduled through a case update, the system can create a hearing row, add a timeline record, and create "
        "notifications so that different modules stay synchronized through shared database state."
    )

    add_section_heading(document, "Data Integrity and Security Considerations:")
    add_body(
        document,
        "The schema uses primary keys, unique constraints, and foreign keys to maintain integrity. Examples include "
        "unique case numbers, unique Aadhaar numbers for citizens, unique Bar Council IDs for advocates, and foreign "
        "keys linking hearings, documents, notes, and tasks to specific cases.\n\n"
        "Authentication is protected through JWT tokens for session access, password hashing for password-based roles, "
        "and time-bound OTP verification for citizens. The document upload module validates file extensions and stores "
        "upload metadata, while role-based route checks prevent unauthorized creation or deletion of records."
    )
    document.add_page_break()


def chapter_five(document: Document, assets: dict[str, Path]) -> None:
    add_chapter_heading(document, 5, "Implementation")

    add_section_heading(document, "Back End via Flask and SQLAlchemy:")
    add_body(
        document,
        "The backend is built with Flask and organized through blueprints for authentication, cases, hearings, "
        "documents, tasks, notes, notifications, messages, courtrooms, and analytics. The application factory "
        "initializes CORS, JWT management, mail support, password hashing, and SQLAlchemy before registering all "
        "API modules.\n\n"
        "The database implementation centers on SQLAlchemy models for users, cases, hearings, case timelines, "
        "documents, tasks, case notes, notifications, messages, courtrooms, and OTP codes. Each model exposes "
        "structured serialization methods so that the React frontend can consume normalized JSON responses.\n\n"
        "Business logic in route handlers includes automatic case number generation, role-based access checks, "
        "case search, QR lookup, hearing calendar endpoints, notification creation, and direct messaging."
    )

    add_section_heading(document, "Front End via React and Vite:")
    add_body(
        document,
        "The frontend uses BrowserRouter and protected routes to separate public, advocate, and court dashboards. "
        "Shared context providers handle authentication state, theme state, and toast notifications, while a central "
        "API service module manages REST communication with the backend.\n\n"
        "Dashboard pages and shared components provide a practical interface for reviewing case details, opening "
        "notifications, managing conversations, tracking hearings, and switching between role-specific workspaces. "
        "The layout encourages quick movement between overview metrics and detailed case actions."
    )
    add_figure(document, assets["dashboards"], "Fig. 5.1 Role-Based Dashboard Overview")

    add_section_heading(document, "Operational Workflow:")
    add_body(
        document,
        "A typical workflow starts with user authentication. Citizens log in using Aadhaar-based OTP verification, "
        "advocates use Bar Council credentials, and court admins use administrator credentials. Once authenticated, "
        "the user receives only the views and records relevant to that role.\n\n"
        "Case creation, hearing scheduling, document uploads, timeline updates, notifications, and analytics are all "
        "tied back to the same relational data model. This reduces duplication and allows changes in one module to "
        "be reflected across the rest of the application."
    )
    add_figure(document, assets["workflow"], "Fig. 5.2 Case Processing Workflow")
    document.add_page_break()


def chapter_six(document: Document) -> None:
    add_chapter_heading(document, 6, "Result and Discussion")

    add_section_heading(document, "System Functionality Evaluation:")
    add_list(document, FUNCTIONAL_RESULTS)

    add_section_heading(document, "Performance Evaluation:")
    add_body(
        document,
        "The project demonstrates strong functional alignment between database design and application behavior. "
        "Role-aware filtering prevents unnecessary data exposure, while relational links ensure that cases, hearings, "
        "documents, tasks, and notifications stay connected.\n\n"
        "A review verification script in the backend validates key flows such as citizen OTP login, advocate login, "
        "court login, role-based case visibility, and propagation of status updates back to the citizen view. This "
        "provides practical evidence that the major workflow paths are working correctly.\n\n"
        "Formal load testing has not been carried out, so the current evaluation is limited to functional correctness "
        "and development-scale usage. The OTP flow is also intentionally simplified for demonstration, because the "
        "generated OTP is returned by the API during review mode."
    )

    add_section_heading(document, "User Feedback and Adoption:")
    add_body(
        document,
        "As an academic project, the system is best interpreted as a realistic prototype for legal workflow "
        "digitization. The main usability strengths are transparency for citizens, dedicated work areas for advocates, "
        "and centralized monitoring for court staff.\n\n"
        "The modular design also makes the project easy to demonstrate and extend. Because each major function is "
        "mapped to a dedicated table and route group, future refinements can be introduced without reworking the "
        "entire system."
    )

    add_section_heading(document, "Future Directions and Enhancements:")
    add_list(document, FUTURE_ENHANCEMENTS)
    document.add_page_break()


def chapter_seven(document: Document) -> None:
    add_chapter_heading(document, 7, "Conclusion")
    add_body(
        document,
        "The Legal Case Management System shows how a well-designed relational database can power a multi-role web "
        "application in a legally relevant domain. By centralizing case records, hearing schedules, documents, notes, "
        "tasks, notifications, and messages, the system improves traceability and reduces fragmentation.\n\n"
        "The project successfully combines React, Flask, SQLAlchemy, and MySQL to deliver a practical academic "
        "prototype that emphasizes integrity, transparency, and modular growth. With stronger production hardening "
        "and deployment features, the same foundation can evolve into a more complete court process platform."
    )
    document.add_page_break()


def chapter_eight(document: Document) -> None:
    add_chapter_heading(document, 8, "References")
    add_list(document, REFERENCES)


def build_report(template_path: Path, output_path: Path) -> Path:
    assets = build_assets()
    document = Document(str(template_path))
    clear_document(document)

    add_cover_page(document)
    add_certificate(document)
    add_declaration(document)
    add_abstract(document)
    add_front_matter(document)

    chapter_one(document)
    chapter_two(document)
    chapter_three(document, assets)
    chapter_four(document)
    chapter_five(document, assets)
    chapter_six(document)
    chapter_seven(document)
    chapter_eight(document)

    document.save(str(output_path))
    return output_path


def main() -> None:
    parser = argparse.ArgumentParser(description="Generate the LCMS DBMS report from the supplied Word template.")
    parser.add_argument("--template", type=Path, default=DEFAULT_TEMPLATE, help="Path to the source .docx template.")
    parser.add_argument("--output", type=Path, default=DEFAULT_OUTPUT, help="Path where the generated .docx will be saved.")
    args = parser.parse_args()

    if not args.template.exists():
        raise FileNotFoundError(f"Template not found: {args.template}")

    args.output.parent.mkdir(parents=True, exist_ok=True)
    output = build_report(args.template, args.output)
    print(output)


if __name__ == "__main__":
    main()
